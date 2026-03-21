"""
Document text extraction — supports PDF, DOCX, XLSX, TXT.
Returns plain text for AI analysis.
"""
import io
import logging

logger = logging.getLogger(__name__)


def extract_text(file_path: str, mime_type: str = '') -> tuple[str, int]:
    """
    Extract plain text from a document file.
    Returns (text, page_count).
    """
    ext = file_path.lower().split('.')[-1]

    if ext == 'pdf' or 'pdf' in mime_type:
        return _extract_pdf(file_path)
    elif ext in ('docx', 'doc') or 'word' in mime_type:
        return _extract_docx(file_path)
    elif ext in ('xlsx', 'xls') or 'excel' in mime_type or 'spreadsheet' in mime_type:
        return _extract_xlsx(file_path)
    elif ext == 'txt' or 'text/plain' in mime_type:
        return _extract_txt(file_path)
    else:
        logger.warning(f"Unsupported file type: {ext} / {mime_type}")
        return '', 0


def _extract_pdf(path: str) -> tuple[str, int]:
    try:
        import PyPDF2
        text_parts = []
        with open(path, 'rb') as f:
            reader = PyPDF2.PdfReader(f)
            page_count = len(reader.pages)
            for page in reader.pages:
                try:
                    text_parts.append(page.extract_text() or '')
                except Exception:
                    pass
        return '\n\n'.join(text_parts), page_count
    except ImportError:
        logger.error("PyPDF2 not installed")
        return '', 0
    except Exception as e:
        logger.error(f"PDF extraction error: {e}")
        return '', 0


def _extract_docx(path: str) -> tuple[str, int]:
    try:
        from docx import Document
        doc = Document(path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        # Also extract tables
        for table in doc.tables:
            for row in table.rows:
                row_text = ' | '.join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    paragraphs.append(row_text)
        return '\n'.join(paragraphs), 1
    except Exception as e:
        logger.error(f"DOCX extraction error: {e}")
        return '', 0


def _extract_xlsx(path: str) -> tuple[str, int]:
    try:
        import openpyxl
        wb = openpyxl.load_workbook(path, read_only=True, data_only=True)
        parts = []
        for sheet in wb.worksheets:
            parts.append(f"[Sheet: {sheet.title}]")
            for row in sheet.iter_rows(values_only=True):
                row_text = ' | '.join(str(v) for v in row if v is not None)
                if row_text.strip():
                    parts.append(row_text)
        return '\n'.join(parts), len(wb.worksheets)
    except Exception as e:
        logger.error(f"XLSX extraction error: {e}")
        return '', 0


def _extract_txt(path: str) -> tuple[str, int]:
    try:
        with open(path, 'r', encoding='utf-8', errors='ignore') as f:
            text = f.read()
        return text, 1
    except Exception as e:
        logger.error(f"TXT extraction error: {e}")
        return '', 0
